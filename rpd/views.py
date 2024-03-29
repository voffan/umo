from datetime import datetime

from django.http import StreamingHttpResponse, HttpResponseRedirect
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.contrib.auth.models import Group as auth_groups
from django.shortcuts import render, redirect
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Border, Alignment, Protection, Font, Side
from transliterate import translit

import rpd.forms as forms
import io
from parsel import Selector
from django.forms import formset_factory
from django.shortcuts import get_object_or_404


from rpd.models import RPDDiscipline, DisciplineResult, Basement, RPDDisciplineContentHours, RPDDisciplineContent, \
    PracticeDescription, DisciplineRating, MarkScale, FOS, Bibliography
from umo.models import (Teacher, Group, GroupList, Synch, Year, EduProgram, Student, Discipline, CheckPoint, Control,
                        DisciplineDetails, BRSpoints, EduPeriod, ExamMarks, Exam, Course, Person, Competency)


def index(request):
    if auth_groups.objects.get(name='teacher') in request.user.groups.all():
        return redirect('disciplines:mysubjects')
    else:
        return redirect('disciplines:disciplines_list')


# Create your views here.
class RPDList(ListView):
    template_name = 'rpd_list.html'
    context_object_name = 'list'

    def get_queryset(self):
        return RPDDiscipline.objects.all()


def rpd_create(request, rpddiscipline_id):
    discipline = RPDDiscipline.objects.filter(id=rpddiscipline_id).first()
    if request.method == 'POST':
        form = forms.RPDProgram(request.POST)
        if form.is_valid():
            rpd = RPDDiscipline.objects.filter(id=rpddiscipline_id).first()
            rpd.goal = form.cleaned_data['goal']
            rpd.abstract = form.cleaned_data['abstract']
            rpd.language = form.cleaned_data['language']
            rpd.education_methodology = form.cleaned_data['education_methodology']
            rpd.count_e_method_support = form.cleaned_data['count_e_method_support']
            rpd.methodological_instructions = form.cleaned_data['methodological_instructions']
            rpd.fos_fond = form.cleaned_data['fos_fond']
            rpd.fos_methodology = form.cleaned_data['fos_methodology']
            rpd.scaling_methodology = form.cleaned_data['scaling_methodology']
            rpd.web_resource = form.cleaned_data['web_resource']
            rpd.material = form.cleaned_data['material']
            rpd.it = form.cleaned_data['it']
            rpd.software = form.cleaned_data['software']
            rpd.iss = form.cleaned_data['iss']
            rpd.save()
        results = forms.ResultsSet(request.POST, prefix='results')
        if results.is_valid():
            for form in results:
                result = form.save(commit=False)
                result.rpd = rpd
                result.save()
        basement = forms.BasementSet(request.POST, prefix='basement')
        if basement.is_valid():
            for form in basement:
                base = form.save(commit=False)
                base.rpd = rpd
                base.save()
        hours_distribution = forms.HoursDistributionSet(request.POST, prefix='hours_distribution')
        if hours_distribution.is_valid():
            for form in hours_distribution:
                hour_dis = form.save(commit=False)
                hour_dis.rpd = rpd
                hour_dis.save()
        theme = forms.ThemeSet(request.POST, prefix='theme')
        if theme.is_valid():
            for form in theme:
                name = form.save(commit=False)
                name.rpd = rpd
                name.save()
        srs_content = forms.SRSContentSet(request.POST, prefix='srs_content')
        if srs_content.is_valid():
            for form in srs_content:
                srs = form.save(commit=False)
                srs.rpd = rpd
                srs.save()
        labs_content = forms.LabContentSet(request.POST, prefix='labs_content')
        if labs_content.is_valid():
            for form in labs_content:
                labs = form.save(commit=False)
                labs.rpd = rpd
                labs.save()
        disc_rating = forms.DiscRatingSet(request.POST, prefix='disc_rating')
        if disc_rating.is_valid():
            for form in disc_rating:
                disc = form.save(commit=False)
                disc.rpd = rpd
                disc.save()
        mark_scale = forms.MarkScaleSet(request.POST, prefix='mark_scale')
        if mark_scale.is_valid():
            for form in mark_scale:
                mark = form.save(commit=False)
                mark.rpd = rpd
                mark.save()
        fos_table = forms.FosTableSet(request.POST, prefix='fos_table')
        if fos_table.is_valid():
            for form in fos_table:
                f_t = form.save(commit=False)
                f_t.rpd = rpd
                f_t.save()
        bibl = forms.BibliographySet(request.POST, prefix='bibl')
        if bibl.is_valid():
            for form in bibl:
                main = form.save(commit=False)
                main.rpd = rpd
                main.save()
        return redirect('rpds:rpd', discipline.id)
    else:
        data = {'goal': discipline.goal,
                'abstract': discipline.abstract,
                'language': discipline.language,
                'education_methodology': discipline.education_methodology,
                'count_e_method_support': discipline.count_e_method_support,
                'methodological_instructions': discipline.methodological_instructions,
                'fos_fond': discipline.fos_fond,
                'fos_methodology': discipline.fos_methodology,
                'scaling_methodology': discipline.scaling_methodology,
                'web_resource': discipline.web_resource,
                'material': discipline.material,
                'it': discipline.it,
                'software': discipline.software,
                'iss': discipline.iss
                }
        form = forms.RPDProgram(data=data)
        results = forms.ResultsSet(queryset=DisciplineResult.objects.filter(rpd__id=rpddiscipline_id), prefix='results', )
        basement = forms.BasementSet(queryset=Basement.objects.filter(discipline__id=rpddiscipline_id), prefix='basement')
        hours_distribution = forms.HoursDistributionSet(queryset=RPDDisciplineContentHours.objects.filter(content__rpd__id=rpddiscipline_id), prefix='hours_distribution')
        theme = forms.ThemeSet(queryset=RPDDisciplineContent.objects.filter(rpd__id=rpddiscipline_id), prefix='theme')
        srs_content = forms.SRSContentSet(queryset=PracticeDescription.objects.filter(rpd__id=rpddiscipline_id), prefix='srs_content')
        labs_content = forms.LabContentSet(queryset=PracticeDescription.objects.filter(rpd__id=rpddiscipline_id), prefix='labs_content')
        disc_rating = forms.DiscRatingSet(queryset=DisciplineRating.objects.filter(rpd__id=rpddiscipline_id), prefix='disc_rating')
        mark_scale = forms.MarkScaleSet(queryset=MarkScale.objects.filter(rpd__id=rpddiscipline_id), prefix='mark_scale')
        fos_table = forms.FosTableSet(queryset=FOS.objects.filter(rpd__id=rpddiscipline_id), prefix='fos_table')
        bibl = forms.BibliographySet(queryset=Bibliography.objects.filter(rpd__id=rpddiscipline_id, is_main=True), prefix='bibl')
    return render(request, 'rpd.html', context={'discipline': discipline, 'form': form, 'results': results, 'basement': basement,
                                                    'hours_distribution': hours_distribution, 'theme': theme, 'srs_content': srs_content,
                                                    'labs_content': labs_content, 'disc_rating': disc_rating, 'mark_scale': mark_scale,
                                                    'fos_table': fos_table, 'bibl': bibl})


def export_docx(request, rpddiscipline_id):
    if request.method == 'GET':
        discipline = RPDDiscipline.objects.filter(id=rpddiscipline_id).first()
        doc = create_docx(discipline)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        file_name = translit(discipline.Name, 'ru', reversed=True)
        response = StreamingHttpResponse(streaming_content=buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=rpd_' + file_name + '.docx'
        response["Content-Encoding"] = 'UTF-8'
        return response


def create_docx(discipline):
    document = Document()
    title = document.add_heading('1. Аннотация к рабочей программе дисциплины', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = document.add_heading(discipline.code, 0)
    t.add_run(discipline.Name)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    z = document.add_heading('Трудоемкость __з.е.', 0)
    z.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading('1.1. Цель освоения и краткое содержание дисциплины', level=1)
    goal = document.add_paragraph('Цель освоения:  ')
    goal.add_run(discipline.goal)
    abstract = document.add_paragraph('Краткое содержание дисциплины:  ')
    abstract.add_run(discipline.abstract)
    document.add_heading('1.2. Перечень планируемых результатов обучения по дисциплине, соотнесенных с планируемыми результатами освоения образовательной программы', level=1)
    table_results = document.add_table(rows=1, cols=5)
    table_results.style = 'Table Grid'
    hdr_cells = table_results.rows[0].cells
    hdr_cells[0].text = 'Наименование категории (группы) компетенций'
    hdr_cells[1].text = 'Планируемые результаты освоения программы (код и содержание компетенции)'
    hdr_cells[2].text = 'Индикаторы достижения компетенций'
    hdr_cells[3].text = 'Планируемые результаты обучения по дисциплине'
    hdr_cells[4].text = 'Оценочные средства'
    results = discipline.disciplineresult_set.order_by('competency__type').all()
    if results.count() > 0:
        comp_type = 0
        for result in results:
            row_cells = table_results.add_row().cells
            if comp_type != result.competency.type:
                row_cells[0].text = Competency.CompetencyType[result.competency.type-1][1]
                comp_type = result.competency.type
            row_cells[1].text = result.competency.short_name + ' ' + result.competency.name
            row_cells[2].text = result.indicator.short_name + ' ' + result.indicator.indicator
            row_cells[3].text = result.skill
            row_cells[4].text = result.judging
    document.add_heading('1.3. Место дисциплины в структуре ОПОП', level=1)
    table_basement = document.add_table(rows=1, cols=6)
    table_basement.style = 'Table Grid'
    hdr_cells = table_basement.rows[0].cells
    hdr_cells[0].text = 'Индексы'
    hdr_cells[1].text = 'Наименование дисциплины (модуля), практики'
    hdr_cells[2].text = 'Семестр изучения'
    hdr_cells[3].text = 'Индексы и наименования учебных дисциплин (модулей), практик'
    hdr_cells[4].text = 'на которые опирается содержание данной дисциплины (модуля)'
    hdr_cells[5].text = 'для которых содержание данной дисциплины (модуля) выступает опорой'
    # basement = discipline.basement_set.order_by('discipline').all()
    # if basement.count() > 0:
    #     d_id = 0
    #     for basement in basement:
    #         row_cells = table_basement.add_row().cells
    #         if d_id != basement.discipline:
    #             row_cells[0].text = Basement.discipline.code
    #             d_id = basement.discipline.code
    #         row_cells[1] = basement.discipline.Name
    #         row_cells[2] = basement.discipline.semester
    #         row_cells[3] = 'Индексы и наименования учебных дисциплин (модулей), практик'
    #         row_cells[4] = basement.discipline
    #         row_cells[5] = basement.base
    document.add_heading('1.4 Язык преподавания: ', level=1)
    #document.add_paragraph(discipline.language)
    title2 = document.add_heading('2. Объем дисциплины в зачетных единицах с указанием количества академических часов, выделенных на контактную работу обучающихся с преподавателем (по видам учебных занятий) и на самостоятельную работу обучающихся', level=1)
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title3 = document.add_heading('3. Содержание дисциплины, структурированное по темам с указанием отведенного на них количества академических часов и видов учебных занятий', level=1)
    title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading('3.1. Распределение часов по темам и видам учебных занятий', level=1)
    table_hours_distribution = document.add_table(rows=1, cols=3)
    table_hours_distribution.style = 'Table Grid'
    hdr_cells = table_hours_distribution.rows[0].cells
    hdr_cells[0].text = 'Тема'
    hdr_cells[1].text = 'Вид учебного занятия'
    hdr_cells[2].text = 'Количеcтво часов'
    #hours_distribution = discipline.rpddisciplinecontenthours_set.order_by('content').all()
    #if hours_distribution.count() > 0:
    #     cont = 0
    #     for hours_distribution in hours_distribution:
    #         row_cells = table_hours_distribution.add_row().cells
    #         if cont != hours_distribution.content:
    #             row_cells[0].text = hours_distribution.content
    #             cont = hours_distribution.competency.type
    #         row_cells[1].text = hours_distribution.hours_type
    #         row_cells[2].text = hours_distribution.hours
    document.add_heading('3.2. Содержание тем программы дисциплины', level=1)
    document.add_heading('3.3. Формы и методы проведения занятий, применяемые учебные технологии', level=1)
    document.add_paragraph(discipline.education_methodology)
    title4 = document.add_heading('4. Перечень учебно-методического обеспечения для самостоятельной работы обучающихся по дисциплине', level=1)
    title4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(discipline.count_e_method_support)
    title_srs = document.add_heading('Содержание СРС', level=3)
    title_srs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_srs = document.add_table(rows=1, cols=4)
    table_srs.style = 'Table Grid'
    hdr_cells = table_srs.rows[0].cells
    hdr_cells[0].text = 'Наименование раздела (темы) дисциплины'
    hdr_cells[1].text = 'Вид СРС'
    hdr_cells[2].text = 'Трудоемкость (в часах)'
    hdr_cells[3].text = 'Формы и методы контроля'
    # srs_content = discipline.practicedescription_set.order_by('practice_type').all()
    # if srs_content.count() > 0:
    #     srs = 0
    #     for srs_content in srs_content:
    #         row_cells = table_hours_distribution.add_row().cells
    #         if srs != srs_content.theme:
    #             row_cells[0].text = srs_content.theme[results.practice_type-1][1]
    #             srs = srs_content.srs_conten.theme
    #         row_cells[1].text = srs_content.practice_type
    #         row_cells[2].text = srs_content.hours
    #         row_cells[3].text = srs_content.control
    title_labs = document.add_heading('Лабораторные работы или лабораторные практикумы', level=3)
    title_labs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_srs = document.add_table(rows=1, cols=4)
    table_srs.style = 'Table Grid'
    hdr_cells = table_srs.rows[0].cells
    hdr_cells[0].text = 'Наименование раздела (темы) дисциплины'
    hdr_cells[1].text = 'Лабораторная работа или лабораторный практикум'
    hdr_cells[2].text = 'Трудоемкость (в часах)'
    hdr_cells[3].text = 'Формы и методы контроля'
    pass
    title5 = document.add_heading('5. Методические указания для обучающихся по освоению дисциплины', level=1)
    title5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(discipline.methodological_instructions)
    document.add_paragraph('Рейтинговый регламент по дисциплине:')
    table_rating = document.add_table(rows=1, cols=5)
    table_rating.style = 'Table Grid'
    hdr_cells = table_rating.rows[0].cells
    hdr_cells[0].text = 'Вид рейтинговой таблицы'
    hdr_cells[1].text = 'Семестр'
    hdr_cells[2].text = 'Вид работы'
    hdr_cells[3].text = 'Макс.балл'
    hdr_cells[4].text = 'Мин.балл'
    pass
    title6 = document.add_heading('6. Фонд оценочных средств для проведения промежуточной аттестации обучающихся по дисциплине', level=1)
    title6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(discipline.fos_fond)
    document.add_heading('6.1. Показатели, критерии и шкала оценивания', level=1)
    table_mark = document.add_table(rows=1, cols=7)
    table_mark.style = 'Table Grid'
    hdr_cells = table_mark.rows[0].cells
    hdr_cells[0].text = 'Коды оцениваемых компетенций'
    hdr_cells[1].text = 'Индикаторы достижения компетенций'
    hdr_cells[2].text = 'Показатель оценивания (по п.1.2.РПД)'

    pass
    document.add_heading('6.2. Примерные контрольные задания (вопросы) для промежуточной аттестации', level=1)
    document.add_paragraph(discipline.fos_methodology)
    table_f = document.add_table(rows=1, cols=5)
    table_f.style = 'Table Grid'
    hdr_cells = table_f.rows[0].cells
    hdr_cells[0].text = 'Коды оцениваемых компетенций'
    hdr_cells[1].text = 'Индикаторы достижения компетенций'
    hdr_cells[2].text = 'Оцениваемый показатель (ЗУВ)'
    hdr_cells[3].text = 'Тема (темы)'
    hdr_cells[4].text = 'Образец типового (тестового или практического) задания (вопроса)'
    pass
    document.add_heading('6.2. Примерные контрольные задания (вопросы) для промежуточной аттестации', level=1)
    document.add_paragraph(discipline.scaling_methodology)
    title7 = document.add_heading('7. Перечень основной и дополнительной учебной литературы, необходимой для освоения дисциплины', level=1)
    title7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_mb = document.add_table(rows=1, cols=5)
    table_mb.style = 'Table Grid'
    hdr_cells = table_mb.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Библиографическая ссылка'
    hdr_cells[2].text = 'Электронная литература'
    hdr_cells[3].text = 'Наличие грифа'
    hdr_cells[4].text = 'Количество экземпляров'
    pass
    title8 = document.add_heading('8. Перечень ресурсов информационно-телекоммуникационной сети «Интернет» (далее сеть-Интернет), необходимых для освоения дисциплины', level=1)
    title8.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(discipline.web_resource)
    title9 = document.add_heading('9. Описание материально-технической базы, необходимой для осуществления образовательного процесса по дисциплине ', level=1)
    title9.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(discipline.material)
    title10 = document.add_heading('10. Перечень информационных технологий, используемых при осуществлении образовательного процесса по дисциплине, включая перечень программного обеспечения и информационных справочных систем ', level=1)
    title10.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading('10.1. Перечень информационных технологий, используемых при осуществлении образовательного процесса по дисциплине', level=1)
    document.add_paragraph(discipline.it)
    document.add_heading('10.2. Перечень программного обеспечения', level=1)
    document.add_paragraph(discipline.software)
    document.add_heading('10.3. Перечень информационных справочных систем', level=1)
    document.add_paragraph(discipline.iss)
    return document
